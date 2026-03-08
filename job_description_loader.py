from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.table import _Cell, Table


def _iter_table_text(table: Table) -> list[str]:
    chunks: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            chunks.extend(_iter_cell_text(cell))
    return chunks


def _iter_cell_text(cell: _Cell) -> list[str]:
    chunks: list[str] = []
    for paragraph in cell.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table in cell.tables:
        chunks.extend(_iter_table_text(table))
    return chunks


def _read_docx_text(path: Path) -> str:
    document = Document(str(path))
    chunks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table in document.tables:
        chunks.extend(_iter_table_text(table))

    return "\n".join(chunks).strip()


def read_job_description(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(f"JD file not found: {path}")

    suffix = path.suffix.lower()
    if suffix == ".txt":
        return path.read_text(encoding="utf-8")
    if suffix == ".docx":
        return _read_docx_text(path)

    raise ValueError("Job Description input must be .txt or .docx.")

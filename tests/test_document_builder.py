from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from document_builder import (
    TemplateValidationError,
    assemble_cv_document,
    discover_template_placeholders,
    extract_docx_text,
    preflight_template,
    write_cover_letters_markdown,
)
from tests.test_support import make_workspace_temp_dir
from workflow_definition import TEMPLATE_SECTION_IDS


def _make_template(path: Path) -> None:
    document = Document()
    document.add_paragraph("Summary: {{section_professional_summary}}")
    document.add_paragraph("Skills: {{section_skills_alignment}}")
    document.add_paragraph("Exp1: {{section_experience_1_oldest}}")
    document.add_paragraph("Exp2: {{section_experience_2_previous}}")
    document.add_paragraph("Exp3: {{section_experience_3_latest}}")
    document.save(path)


def test_preflight_and_assemble_docx_with_normalized_experience_placeholders() -> None:
    tmp_path = make_workspace_temp_dir("docx-preflight")
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "output.docx"
    _make_template(template_path)

    preflight_template(template_path, TEMPLATE_SECTION_IDS)
    assemble_cv_document(
        template_path=template_path,
        output_path=output_path,
        selected_content={
            "section_professional_summary": "SUMMARY OUT",
            "section_skills_alignment": "SKILLS OUT",
            "section_experience_1": "EXP1 OUT",
            "section_experience_2": "EXP2 OUT",
            "section_experience_3": "EXP3 OUT",
        },
    )

    rendered = Document(str(output_path))
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "SUMMARY OUT" in text
    assert "SKILLS OUT" in text
    assert "EXP1 OUT" in text
    assert "EXP2 OUT" in text
    assert "EXP3 OUT" in text


def test_duplicate_normalized_placeholders_fail_fast() -> None:
    tmp_path = make_workspace_temp_dir("docx-duplicate")
    template_path = tmp_path / "template_dup.docx"
    document = Document()
    document.add_paragraph("{{section_experience_1_oldest}}")
    document.add_paragraph("{{section_experience_1_latest}}")
    document.save(template_path)

    with pytest.raises(TemplateValidationError):
        discover_template_placeholders(template_path)


def test_extract_docx_text_reads_rendered_text() -> None:
    tmp_path = make_workspace_temp_dir("docx-text")
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "output.docx"
    _make_template(template_path)

    assemble_cv_document(
        template_path=template_path,
        output_path=output_path,
        selected_content={
            "section_professional_summary": "SUMMARY OUT",
            "section_skills_alignment": "SKILLS OUT",
            "section_experience_1": "EXP1 OUT",
            "section_experience_2": "EXP2 OUT",
            "section_experience_3": "EXP3 OUT",
        },
    )

    extracted = extract_docx_text(output_path)
    assert "SUMMARY OUT" in extracted
    assert "EXP3 OUT" in extracted


def test_write_cover_letters_markdown_exports_selected_and_variations() -> None:
    tmp_path = make_workspace_temp_dir("cover-letters")
    output_path = tmp_path / "cover_letters.md"
    write_cover_letters_markdown(
        output_path,
        selected_content="Hello recruiter",
        variations=[
            {
                "id": "A",
                "score_0_to_100": 95,
                "ai_reasoning": "Best fit",
                "content_for_template": "First option",
            },
            {
                "id": "B",
                "score_0_to_100": 91,
                "ai_reasoning": "Second best",
                "content_for_template": "Second option",
            },
        ],
    )

    text = output_path.read_text(encoding="utf-8")
    assert "# Cover Letters" in text
    assert "## Final Approved Version" in text
    assert "Hello recruiter" in text
    assert "## Variation A" in text
    assert "Score: 95" in text

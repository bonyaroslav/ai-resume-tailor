from __future__ import annotations

import argparse
import asyncio
from pathlib import Path

from docx import Document

import main
from settings import DEFAULT_GEMINI_MODEL
from tests.test_support import make_workspace_temp_dir


def _make_template(path: Path) -> None:
    document = Document()
    document.add_paragraph("Summary: {{section_professional_summary}}")
    document.add_paragraph("Skills: {{section_skills_alignment}}")
    document.add_paragraph("Exp1: {{section_experience_1_oldest}}")
    document.add_paragraph("Exp2: {{section_experience_2_previous}}")
    document.add_paragraph("Exp3: {{section_experience_3_latest}}")
    document.save(path)


def test_offline_end_to_end_run_creates_outputs(monkeypatch: object) -> None:
    tmp_dir = make_workspace_temp_dir("offline-e2e")
    run_dir = tmp_dir / "run"
    run_dir.mkdir(parents=True, exist_ok=False)

    jd_path = tmp_dir / "jd.txt"
    jd_path.write_text(
        "Senior Python engineer role focused on API reliability and data workflows.",
        encoding="utf-8",
    )
    template_path = tmp_dir / "template.docx"
    _make_template(template_path)

    args = argparse.Namespace(
        command="run",
        jd_path=jd_path,
        company="offline-smoke",
        job_title="",
        template_path=template_path,
        model=DEFAULT_GEMINI_MODEL,
        role=None,
        debug=False,
        invalidate_cache=False,
        force_knowledge_reupload=False,
        skills_category_count=None,
    )

    monkeypatch.setenv("ART_OFFLINE_MODE", "1")
    monkeypatch.setenv("ART_AUTO_APPROVE_REVIEW", "1")
    monkeypatch.setenv("ART_TRIAGE_DECISION_MODE", "always_continue")
    monkeypatch.setattr(main, "create_run_directory", lambda _, __, ___=None: run_dir)

    asyncio.run(main._handle_run(args))

    cv_path = run_dir / "tailored_cv.docx"
    cover_path = run_dir / "cover_letters.md"
    audit_path = run_dir / "cv_deep_dive_audit.md"
    checkpoint_path = run_dir / "state_checkpoint.json"
    log_path = run_dir / "run.log"

    assert cv_path.exists()
    assert cover_path.exists()
    assert audit_path.exists()
    assert checkpoint_path.exists()
    assert log_path.exists()

    rendered = Document(str(cv_path))
    output_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "Senior software engineer" in output_text
    assert "- Languages, core stack: Python, SQL, JavaScript" in output_text
    assert "Implemented scalable data processing pipelines" in output_text

    cover_text = cover_path.read_text(encoding="utf-8")
    assert "## Final Approved Version" in cover_text
    assert "## Variation A" in cover_text
    assert "I build production Python services and data workflows" in cover_text

    audit_text = audit_path.read_text(encoding="utf-8")
    assert "# Deep Dive CV Audit" in audit_text
    assert "## Prioritized Fixes" in audit_text

    log_text = log_path.read_text(encoding="utf-8")
    assert "Run finished with status=completed." in log_text

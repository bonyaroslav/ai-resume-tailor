from __future__ import annotations

import asyncio
import json
from pathlib import Path

from docx import Document

import graph_nodes
import main as main_module
from checkpoint import load_checkpoint
from graph_nodes import RuntimeContext
from graph_state import GraphState, create_initial_state
from knowledge_cache import RunScopedKnowledgeCache
from llm_client import LlmGenerationResult, UsageMetadata
from main import _run_graph
from prompt_loader import PromptTemplate
from tests.test_support import make_workspace_temp_dir
from workflow_definition import (
    GENERATION_SECTION_IDS,
    WORKFLOW_SECTION_IDS,
)


def _make_template(path: Path) -> None:
    document = Document()
    document.add_paragraph("Summary: {{section_professional_summary}}")
    document.add_paragraph("Skills: {{section_skills_alignment}}")
    document.add_paragraph("Exp1: {{section_experience_1_oldest}}")
    document.add_paragraph("Exp2: {{section_experience_2_previous}}")
    document.add_paragraph("Exp3: {{section_experience_3_latest}}")
    document.save(path)


def _build_prompt_templates(run_dir: Path) -> dict[str, PromptTemplate]:
    templates: dict[str, PromptTemplate] = {}
    for section_id in WORKFLOW_SECTION_IDS:
        templates[section_id] = PromptTemplate(
            section_id=section_id,
            path=run_dir / f"{section_id}.md",
            body=f"SECTION_ID: {section_id}",
            knowledge_files=[],
        )
    return templates


def _build_runtime_context(run_dir: Path, template_path: Path) -> RuntimeContext:
    job_description_path = run_dir / "job_description.md"
    job_description_path.write_text("Example JD", encoding="utf-8")
    context = RuntimeContext(
        run_dir=run_dir,
        checkpoint_path=run_dir / "state_checkpoint.json",
        template_path=template_path,
        output_cv_path=run_dir / "tailored_cv.docx",
        output_cover_letters_path=run_dir / "cover_letters.md",
        output_audit_path=run_dir / "cv_deep_dive_audit.md",
        company_name="Acme",
        job_description_path=job_description_path,
        job_description="Example JD",
        api_key="test-key",
        model_name="fake-model",
        input_profile="role_engineer",
        prompt_templates=_build_prompt_templates(run_dir),
        debug_mode=False,
        auto_approve_review=False,
        triage_decision_mode="prompt",
    )
    context.knowledge_cache_registry_path = run_dir / "cache-registry.json"
    context.knowledge_cache_ttl_seconds = 3600
    return context


def _fake_response_for_section(section_id: str) -> str:
    if section_id == "triage_job_fit_and_risks":
        return json.dumps(
            {
                "triage_result": {
                    "verdict": "APPLY",
                    "decision_score_0_to_100": 85,
                    "confidence_0_to_100": 80,
                    "summary": "Strong alignment for this role.",
                    "raw_subscores": {
                        "technical_fit_0_to_35": 31,
                        "company_risk_0_to_20": 16,
                        "role_quality_0_to_15": 12,
                        "spain_entity_compat_0_to_20": 14,
                        "evidence_quality_0_to_10": 8,
                    },
                    "top_reasons": [
                        "Strong backend fit",
                        "Good role scope",
                        "Manageable risk profile",
                    ],
                    "key_risks": [
                        {
                            "risk": "Need recruiter confirmation on contract type.",
                            "severity": "medium",
                            "type": "uncertainty",
                            "mitigation": "Ask direct hiring-entity questions early.",
                        }
                    ],
                    "spain_entity_risk": {
                        "status": "UNCLEAR",
                        "confidence_0_to_100": 55,
                        "explanation": "Public policy does not clarify B2B constraints.",
                        "recruiter_questions": [
                            "Can you hire via contractor agreement?",
                            "Which legal entity signs the contract?",
                            "Is payroll employment mandatory in Spain?",
                        ],
                    },
                    "sources": [
                        {
                            "label": "Company careers",
                            "url": "https://example.com/careers",
                            "evidence_grade": "A",
                            "used_for": "Hiring policy",
                        }
                    ],
                    "report_markdown": "### Final Verdict\nProceed with this role.",
                }
            }
        )

    if section_id == "section_skills_alignment":
        return json.dumps(
            {
                "meta": {
                    "jd_top_keywords": ["python", "sql"],
                    "covered_keywords": ["python", "sql"],
                    "missing_keywords_not_in_matrix": [],
                },
                "variations": [
                    {
                        "id": "A",
                        "score_0_to_100": 95,
                        "ai_reasoning": f"Reason for {section_id}",
                        "categories": [
                            {
                                "category_name": "Languages, core stack",
                                "category_text": "Python, SQL",
                            },
                            {"category_name": "Cloud, infra", "category_text": "AWS"},
                            {
                                "category_name": "Testing, quality",
                                "category_text": "pytest",
                            },
                            {
                                "category_name": "Delivery, tooling",
                                "category_text": "Docker",
                            },
                        ],
                    },
                    {
                        "id": "B",
                        "score_0_to_100": 91,
                        "ai_reasoning": "Fallback",
                        "categories": [
                            {
                                "category_name": "Languages, core stack",
                                "category_text": "Python",
                            },
                            {"category_name": "Cloud, infra", "category_text": "AWS"},
                            {
                                "category_name": "Testing, quality",
                                "category_text": "pytest",
                            },
                            {
                                "category_name": "Delivery, tooling",
                                "category_text": "Docker",
                            },
                        ],
                    },
                ],
            }
        )

    if section_id == "doc_cover_letter":
        return json.dumps(
            {
                "variations": [
                    {
                        "id": "A",
                        "score_0_to_100": 96,
                        "ai_reasoning": "Best fit",
                        "content_for_template": "Approved content for doc_cover_letter",
                    },
                    {
                        "id": "B",
                        "score_0_to_100": 92,
                        "ai_reasoning": "Strong alternative",
                        "content_for_template": "Alternative B for doc_cover_letter",
                    },
                    {
                        "id": "C",
                        "score_0_to_100": 88,
                        "ai_reasoning": "Balanced option",
                        "content_for_template": "Alternative C for doc_cover_letter",
                    },
                    {
                        "id": "D",
                        "score_0_to_100": 84,
                        "ai_reasoning": "Shortest option",
                        "content_for_template": "Alternative D for doc_cover_letter",
                    },
                ]
            }
        )

    if section_id == "audit_cv_deep_dive":
        return (
            "# Deep Dive CV Audit\n\n"
            "## Executive Summary\nStrong fit with under-signalled evidence.\n\n"
            "## ATS Match Rate\nEstimated match: 84%\n\n"
            "## Keyword Gap Analysis\n- .NET 8: under-represented\n\n"
            "## Hiring Manager Read\nStrongest reason: relevant backend depth.\n\n"
            "## Section-by-Section Critique\n### Professional Summary\nNeeds sharper positioning.\n\n"
            "## Evidence Gaps\n- Metrics are too thin in recent experience.\n\n"
            "## Prioritized Fixes\n1. Strengthen summary fit.\n\n"
            "## Rewrite Directions\nBefore: generic summary\nBetter direction: role-aligned summary.\n\n"
            "## Final Verdict\nCompetitive, but needs stronger signaling.\n"
        )
    envelope = {
        "variations": [
            {
                "id": "A",
                "score_0_to_100": 5,
                "ai_reasoning": f"Reason for {section_id}",
                "content_for_template": f"Approved content for {section_id}",
            },
            {
                "id": "B",
                "score_0_to_100": 3,
                "ai_reasoning": "Fallback",
                "content_for_template": f"Fallback content for {section_id}",
            },
        ]
    }
    return json.dumps(envelope)


def _extract_section_id_from_prompt(prompt: str) -> str:
    marker = "SECTION_ID: "
    for line in prompt.splitlines():
        if line.startswith(marker):
            return line[len(marker) :].strip()
    raise AssertionError(f"Missing section marker in prompt: {prompt}")


def _result(text: str, *, cached_tokens: int | None = None) -> LlmGenerationResult:
    return LlmGenerationResult(
        text=text,
        usage_metadata=UsageMetadata(
            prompt_token_count=10,
            cached_content_token_count=cached_tokens,
            candidates_token_count=5,
            thoughts_token_count=1,
            total_token_count=16,
        ),
    )


def test_run_graph_completes_with_mocked_llm_and_review_choices(
    monkeypatch: object,
) -> None:
    run_dir = make_workspace_temp_dir("integration-mocked-graph")
    template_path = run_dir / "template.docx"
    _make_template(template_path)
    context = _build_runtime_context(run_dir, template_path)
    state: GraphState = create_initial_state("integration-run-1")

    async def fake_generate_with_gemini(
        prompt: str,
        api_key: str,
        model: str,
        section_id: str | None = None,
        cached_content_name: str | None = None,
        skills_category_count: int = 4,
    ) -> LlmGenerationResult:
        assert api_key == "test-key"
        assert model == "fake-model"
        assert cached_content_name is None
        assert skills_category_count == 4
        resolved_section_id = section_id or _extract_section_id_from_prompt(prompt)
        return _result(_fake_response_for_section(resolved_section_id))

    review_inputs = ["continue"]
    for _ in GENERATION_SECTION_IDS:
        review_inputs.extend(["choose", "A"])
    response_iter = iter(review_inputs)

    def fake_input(_: str = "") -> str:
        try:
            return next(response_iter)
        except StopIteration as exc:
            raise AssertionError("Review requested more inputs than expected.") from exc

    monkeypatch.setattr(graph_nodes, "generate_with_gemini", fake_generate_with_gemini)
    monkeypatch.setattr("builtins.input", fake_input)

    final_state = asyncio.run(_run_graph(state, context))

    assert final_state.status == "completed"
    assert context.output_cv_path.exists()
    assert context.output_cover_letters_path.exists()
    assert context.output_audit_path.exists()
    checkpoint_state = load_checkpoint(context.checkpoint_path)
    assert checkpoint_state.status == "completed"
    assert checkpoint_state.section_states["section_skills_alignment"].ai_outputs
    assert checkpoint_state.section_states["section_skills_alignment"].ai_outputs[
        0
    ].parsed_payload["meta"]["covered_keywords"] == ["python", "sql"]

    rendered = Document(str(context.output_cv_path))
    output_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "Approved content for section_professional_summary" in output_text
    assert (
        "Skills\n- Languages, core stack: Python, SQL\n- Cloud, infra: AWS\n- Testing, quality: pytest\n- Delivery, tooling: Docker"
        in output_text
    )
    assert "Approved content for section_experience_1" in output_text
    assert "Approved content for section_experience_2" in output_text
    assert "Approved content for section_experience_3" in output_text

    cover_letters = context.output_cover_letters_path.read_text(encoding="utf-8")
    assert "## Final Approved Version" in cover_letters
    assert "Approved content for doc_cover_letter" in cover_letters
    assert "## Variation A" in cover_letters
    audit_text = context.output_audit_path.read_text(encoding="utf-8")
    assert "# Deep Dive CV Audit" in audit_text


def test_run_graph_stops_at_triage_when_user_selects_stop(monkeypatch: object) -> None:
    run_dir = make_workspace_temp_dir("integration-mocked-triage-stop")
    template_path = run_dir / "template.docx"
    _make_template(template_path)
    context = _build_runtime_context(run_dir, template_path)
    state: GraphState = create_initial_state("integration-run-2")

    async def fake_generate_with_gemini(
        prompt: str,
        api_key: str,
        model: str,
        section_id: str | None = None,
        cached_content_name: str | None = None,
        skills_category_count: int = 4,
    ) -> LlmGenerationResult:
        assert api_key == "test-key"
        assert model == "fake-model"
        assert cached_content_name is None
        assert skills_category_count == 4
        resolved_section_id = section_id or _extract_section_id_from_prompt(prompt)
        if resolved_section_id == "triage_job_fit_and_risks":
            return _result(
                json.dumps(
                    {
                        "triage_result": {
                            "verdict": "AVOID",
                            "decision_score_0_to_100": 22,
                            "confidence_0_to_100": 78,
                            "summary": "High legal and role mismatch risk.",
                            "raw_subscores": {
                                "technical_fit_0_to_35": 15,
                                "company_risk_0_to_20": 4,
                                "role_quality_0_to_15": 5,
                                "spain_entity_compat_0_to_20": 1,
                                "evidence_quality_0_to_10": 7,
                            },
                            "top_reasons": ["Reason 1", "Reason 2", "Reason 3"],
                            "key_risks": [
                                {
                                    "risk": "Likely legal blocker for Spain setup.",
                                    "severity": "high",
                                    "type": "legal_blocker",
                                    "mitigation": "Do not proceed without explicit B2B path.",
                                }
                            ],
                            "spain_entity_risk": {
                                "status": "YES",
                                "confidence_0_to_100": 85,
                                "explanation": "Evidence points to local employment requirement.",
                                "recruiter_questions": ["Q1", "Q2", "Q3"],
                            },
                            "sources": [
                                {
                                    "label": "Policy",
                                    "url": "https://example.com/policy",
                                    "evidence_grade": "A",
                                    "used_for": "Employment constraints",
                                }
                            ],
                            "report_markdown": "### Final Verdict\nAvoid.",
                        }
                    }
                )
            )
        return _result(_fake_response_for_section(resolved_section_id))

    monkeypatch.setattr(graph_nodes, "generate_with_gemini", fake_generate_with_gemini)
    monkeypatch.setattr("builtins.input", lambda _: "stop")

    final_state = asyncio.run(_run_graph(state, context))

    assert final_state.status == "completed"
    assert final_state.current_node == "triage_stop"
    assert not context.output_cv_path.exists()
    assert not context.output_cover_letters_path.exists()
    assert not context.output_audit_path.exists()


def test_run_graph_always_continue_ignores_avoid_triage(monkeypatch: object) -> None:
    run_dir = make_workspace_temp_dir("integration-mocked-triage-always-continue")
    template_path = run_dir / "template.docx"
    _make_template(template_path)
    context = _build_runtime_context(run_dir, template_path)
    context.triage_decision_mode = "always_continue"
    context.auto_approve_review = True
    state: GraphState = create_initial_state("integration-run-4")

    async def fake_generate_with_gemini(
        prompt: str,
        api_key: str,
        model: str,
        section_id: str | None = None,
        cached_content_name: str | None = None,
        skills_category_count: int = 4,
    ) -> LlmGenerationResult:
        resolved_section_id = section_id or _extract_section_id_from_prompt(prompt)
        if resolved_section_id == "triage_job_fit_and_risks":
            return _result(
                json.dumps(
                    {
                        "triage_result": {
                            "verdict": "AVOID",
                            "decision_score_0_to_100": 20,
                            "confidence_0_to_100": 80,
                            "summary": "Avoid.",
                            "raw_subscores": {
                                "technical_fit_0_to_35": 10,
                                "company_risk_0_to_20": 3,
                                "role_quality_0_to_15": 3,
                                "spain_entity_compat_0_to_20": 1,
                                "evidence_quality_0_to_10": 3,
                            },
                            "top_reasons": ["Reason 1", "Reason 2", "Reason 3"],
                            "key_risks": [
                                {
                                    "risk": "High risk",
                                    "severity": "high",
                                    "type": "legal_blocker",
                                    "mitigation": "Mitigate",
                                }
                            ],
                            "spain_entity_risk": {
                                "status": "YES",
                                "confidence_0_to_100": 90,
                                "explanation": "Blocked.",
                                "recruiter_questions": ["Q1", "Q2", "Q3"],
                            },
                            "sources": [
                                {
                                    "label": "Policy",
                                    "url": "https://example.com/policy",
                                    "evidence_grade": "A",
                                    "used_for": "Risk",
                                }
                            ],
                            "report_markdown": "### Final Verdict\nAvoid.",
                        }
                    }
                )
            )
        return _result(_fake_response_for_section(resolved_section_id))

    monkeypatch.setattr(graph_nodes, "generate_with_gemini", fake_generate_with_gemini)
    monkeypatch.setattr(
        "builtins.input",
        lambda _: (_ for _ in ()).throw(AssertionError("input should not be called")),
    )

    final_state = asyncio.run(_run_graph(state, context))

    assert final_state.status == "completed"
    assert final_state.current_node == "completed"
    assert context.output_cv_path.exists()
    assert context.output_cover_letters_path.exists()
    assert context.output_audit_path.exists()


def test_run_graph_passes_cached_content_and_skips_inline_knowledge(
    monkeypatch: object,
) -> None:
    run_dir = make_workspace_temp_dir("integration-mocked-cache")
    template_path = run_dir / "template.docx"
    _make_template(template_path)
    context = _build_runtime_context(run_dir, template_path)
    context.use_role_wide_knowledge_cache = True
    context.cached_content_name = "cachedContents/abc123"
    knowledge_file = run_dir / "knowledge.md"
    knowledge_file.write_text("inline context to skip", encoding="utf-8")
    context.prompt_templates["section_professional_summary"] = PromptTemplate(
        section_id="section_professional_summary",
        path=run_dir / "section_professional_summary.md",
        body="SECTION_ID: section_professional_summary",
        knowledge_files=[knowledge_file],
    )
    state: GraphState = create_initial_state("integration-run-3")
    prompts_seen: list[str] = []
    cached_names: list[str | None] = []

    async def fake_generate_with_gemini(
        prompt: str,
        api_key: str,
        model: str,
        section_id: str | None = None,
        cached_content_name: str | None = None,
        skills_category_count: int = 4,
    ) -> LlmGenerationResult:
        prompts_seen.append(prompt)
        assert skills_category_count == 4
        cached_names.append(cached_content_name)
        resolved_section_id = section_id or _extract_section_id_from_prompt(prompt)
        return _result(_fake_response_for_section(resolved_section_id), cached_tokens=9)

    review_inputs = ["continue"]
    for _ in GENERATION_SECTION_IDS:
        review_inputs.extend(["choose", "A"])
    response_iter = iter(review_inputs)

    def fake_input(_: str = "") -> str:
        return next(response_iter)

    monkeypatch.setattr(graph_nodes, "generate_with_gemini", fake_generate_with_gemini)
    monkeypatch.setattr("builtins.input", fake_input)

    final_state = asyncio.run(_run_graph(state, context))

    assert final_state.status == "completed"
    assert all(name == "cachedContents/abc123" for name in cached_names)
    assert all("inline context to skip" not in prompt for prompt in prompts_seen)
    assert all(
        "- `job_description.md` - Source of truth" in prompt for prompt in prompts_seen
    )


def test_run_graph_prepares_role_wide_knowledge_cache_before_generation(
    monkeypatch: object,
) -> None:
    run_dir = make_workspace_temp_dir("integration-mocked-cache-prepare")
    template_path = run_dir / "template.docx"
    _make_template(template_path)
    context = _build_runtime_context(run_dir, template_path)
    context.use_role_wide_knowledge_cache = True
    context.auto_approve_review = True
    context.triage_decision_mode = "always_continue"
    state: GraphState = create_initial_state("integration-run-5")
    prepared_cache_calls: list[dict[str, str]] = []
    cached_names: list[str | None] = []

    def fake_prepare_run_scoped_knowledge_cache(
        **kwargs: object,
    ) -> RunScopedKnowledgeCache:
        prepared_cache_calls.append(
            {
                "run_id": str(kwargs["run_id"]),
                "input_profile": str(kwargs["input_profile"]),
                "model_name": str(kwargs["model_name"]),
            }
        )
        return RunScopedKnowledgeCache(
            remote_cache_name="cachedContents/prepared-123",
            expires_at="2026-03-16T15:09:29Z",
            stable_fingerprint="fingerprint-123",
            job_description_sha256="jd-sha-123",
        )

    async def fake_generate_with_gemini(
        prompt: str,
        api_key: str,
        model: str,
        section_id: str | None = None,
        cached_content_name: str | None = None,
        skills_category_count: int = 4,
    ) -> LlmGenerationResult:
        assert api_key == "test-key"
        assert model == "fake-model"
        assert skills_category_count == 4
        cached_names.append(cached_content_name)
        resolved_section_id = section_id or _extract_section_id_from_prompt(prompt)
        return _result(_fake_response_for_section(resolved_section_id), cached_tokens=9)

    monkeypatch.setattr(
        main_module,
        "prepare_run_scoped_knowledge_cache",
        fake_prepare_run_scoped_knowledge_cache,
    )
    monkeypatch.setattr(graph_nodes, "generate_with_gemini", fake_generate_with_gemini)
    monkeypatch.setattr(
        "builtins.input",
        lambda _: (_ for _ in ()).throw(AssertionError("input should not be called")),
    )

    final_state = asyncio.run(_run_graph(state, context))

    assert final_state.status == "completed"
    assert prepared_cache_calls == [
        {
            "run_id": "integration-run-5",
            "input_profile": "role_engineer",
            "model_name": "fake-model",
        }
    ]
    assert context.cached_content_name == "cachedContents/prepared-123"
    assert all(name == "cachedContents/prepared-123" for name in cached_names)

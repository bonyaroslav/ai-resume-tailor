from __future__ import annotations

import re
from collections.abc import Iterable
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.table import _Cell, Table

from section_ids import is_experience_section, normalize_section_id

PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


class TemplateValidationError(ValueError):
    pass


def _iter_paragraphs(document: DocumentType) -> Iterable:
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table) -> Iterable:
    for row in table.rows:
        for cell in row.cells:
            yield from _iter_cell_paragraphs(cell)


def _iter_cell_paragraphs(cell: _Cell) -> Iterable:
    for paragraph in cell.paragraphs:
        yield paragraph
    for nested_table in cell.tables:
        yield from _iter_table_paragraphs(nested_table)


def discover_template_placeholders(template_path: Path) -> dict[str, str]:
    document = Document(str(template_path))
    placeholders: dict[str, str] = {}

    for paragraph in _iter_paragraphs(document):
        matches = PLACEHOLDER_PATTERN.findall(paragraph.text)
        for raw_name in matches:
            canonical_name = normalize_section_id(raw_name.strip())
            if canonical_name in placeholders:
                raise TemplateValidationError(
                    f"Duplicate template placeholders normalize to section_id '{canonical_name}'."
                )
            placeholders[canonical_name] = raw_name.strip()
    return placeholders


def preflight_template(
    template_path: Path, required_section_ids: Iterable[str]
) -> None:
    placeholders = discover_template_placeholders(template_path)
    missing = [
        section_id
        for section_id in required_section_ids
        if section_id not in placeholders
    ]
    if missing:
        joined = ", ".join(missing)
        raise TemplateValidationError(f"Missing required DOCX placeholders: {joined}")


def _placeholder_regex_for_section(section_id: str) -> re.Pattern[str]:
    escaped = re.escape(section_id)
    if is_experience_section(section_id):
        return re.compile(r"\{\{\s*" + escaped + r"(?:_[^{}]+)?\s*\}\}")
    return re.compile(r"\{\{\s*" + escaped + r"\s*\}\}")


def assemble_cv_document(
    template_path: Path,
    output_path: Path,
    selected_content: dict[str, str],
) -> None:
    document = Document(str(template_path))
    patterns = {
        section_id: _placeholder_regex_for_section(section_id)
        for section_id in selected_content
    }

    for paragraph in _iter_paragraphs(document):
        text = paragraph.text
        replaced = text
        for section_id, content in selected_content.items():
            replaced = patterns[section_id].sub(content, replaced)
        if replaced != text:
            paragraph.text = replaced

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def write_cover_letter(output_path: Path, content: str) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(content.strip() + "\n", encoding="utf-8")
